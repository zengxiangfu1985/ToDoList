#!/usr/bin/env python3
"""Convert ToDoList user manual Markdown to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_rich_text(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label = re.match(r"\[([^\]]+)\]", token)
            if label:
                paragraph.add_run(label.group(1))
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_doc_defaults(doc)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            if not in_code and code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            if level == 1:
                p = doc.add_heading(title, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = header
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    if c < len(table.rows[r + 1].cells):
                        table.rows[r + 1].cells[c].text = value
            doc.add_paragraph()
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_text(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Generated: {docx_path}")


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    if len(sys.argv) >= 3:
        md_path = Path(sys.argv[1])
        docx_path = Path(sys.argv[2])
    else:
        md_path = root / "docs" / "ToDoList用户使用说明.md"
        docx_path = root / "docs" / "ToDoList用户使用说明.docx"
    if not md_path.exists():
        print(f"Missing: {md_path}", file=sys.stderr)
        return 1
    md_to_docx(md_path, docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
